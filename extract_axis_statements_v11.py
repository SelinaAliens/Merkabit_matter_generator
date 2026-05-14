"""Extract every paragraph in Paper 41 v11 that mentions an axis or Re(s)/Im(s).

Output:
  - paragraph index
  - first 80 chars of the surrounding section heading (if any nearby)
  - full paragraph text

This lets us list and review every axis-statement before editing.
"""
from pathlib import Path
from docx import Document

HERE = Path(__file__).resolve().parent
DOC = HERE / "Paper_41_Matter_Generator_v11.docx"

KEYWORDS = [
    "x-axis", "y-axis", "z-axis",
    "X-axis", "Y-axis", "Z-axis",
    "all-T", "all-S", "all-R",
    "Re(s)", "Im(s)", "Re(s)=", "Im(s)=",
    "imaginary axis", "real axis",
    "imaginary part", "real part",
    "substrate clock", "matter emerges",
    "T-direction", "S-direction", "R-direction",
    "Re(s) = 1/2", "critical line",
]


def main():
    doc = Document(str(DOC))
    print("Paper_41_Matter_Generator_v11.docx — axis-statement scan")
    print("=" * 70)
    print("Total paragraphs: %d" % len(doc.paragraphs))
    print()

    last_heading = "(no heading yet)"
    hits = 0
    for i, p in enumerate(doc.paragraphs):
        # Track headings
        try:
            sname = p.style.name if p.style else ""
        except Exception:
            sname = ""
        if sname and sname.startswith("Heading"):
            last_heading = "[%s] %s" % (sname, p.text.strip()[:100])

        text = p.text
        if not text or not text.strip():
            continue

        # Match any keyword
        matched = [k for k in KEYWORDS if k in text]
        if matched:
            hits += 1
            print("=" * 70)
            print("HIT #%d  para_idx=%d  matched=%s" % (hits, i, matched))
            print("  under: %s" % last_heading)
            print()
            # Print full text wrapped
            for line in text.split("\n"):
                print("  | " + line)
            print()

    print("=" * 70)
    print("Total hits: %d" % hits)


if __name__ == "__main__":
    main()
