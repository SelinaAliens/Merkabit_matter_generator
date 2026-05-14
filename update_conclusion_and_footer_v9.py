"""
Update §10 Conclusion and version footer of Paper_41_Matter_Generator_v9.docx
to reflect the §2.9 bipartite-vacuum / Λ-rate addition.

Changes:
  1. Insert new §10 paragraph after the §2.8-dynamic-layer paragraph,
     describing the §2.9 bipartite-vacuum / Λ-rate / two-architectural-rates
     reading.
  2. Update the crown-jewel paragraph to include the bipartite vacuum +
     two architectural rates + dark energy operational reading.
  3. Update version footer: v6 → v9, add v7 / v8 / v9 changelog entries,
     update "four complementary" → "five complementary".
"""
from pathlib import Path
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

HERE = Path(__file__).resolve().parent
DOC = HERE / "Paper_41_Matter_Generator_v9.docx"


def make_para(text, italic=False):
    """Create a paragraph element with optional italic styling."""
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    p.append(pPr)
    r = OxmlElement("w:r")
    if italic:
        rPr = OxmlElement("w:rPr")
        rPr.append(OxmlElement("w:i"))
        r.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    r.append(t)
    p.append(r)
    return p


def replace_paragraph_text(para, new_text):
    """Replace all text in a paragraph with new_text, preserving first run formatting."""
    if not para.runs:
        para.add_run(new_text)
        return
    # Keep first run; clear it; remove other runs
    para.runs[0].text = new_text
    for r in list(para.runs)[1:]:
        r._element.getparent().remove(r._element)


def main():
    print("Opening %s..." % DOC.name)
    doc = Document(str(DOC))

    # ----------------------------------------------------------------
    # STEP 1: Insert new §10 paragraph after the §2.8-dynamic paragraph
    # The §2.8 paragraph starts with "At the dynamic level (§2.8)"
    # The new §2.9 paragraph should go after it, before the RH statement
    # ----------------------------------------------------------------
    new_para_text = (
        "At the bipartite-vacuum level (§2.9, Paper 23 [23]), the operating "
        "geometry sits inside a global ground state where every node carries "
        "a dual-spinor (|u⟩, |v⟩) with |u⟩ ⊥ |v⟩ — the zero-point standing "
        "wave that the matter generator builds against. The architecture "
        "produces TWO simultaneous rates from two combinations of the same "
        "fundamental quantities: η_B = 5α⁴/|2T| ≈ 5.91 × 10⁻¹⁰ per cycle "
        "(matter creation success rate at the operating site, §2.8.3) and "
        "Λ = √(3/2) × exp(−2γ_Berry × dim(E₆) × h / 2π) ≈ 2.87 × 10⁻¹²² per "
        "cycle per node (vacuum monopole formation rate at every node "
        "simultaneously, §2.9.3). The ratio η_B / Λ ≈ 2.6 × 10¹¹² — matter "
        "creation is exponentially more likely than vacuum self-disruption, "
        "but both are exponential suppressions arising from the architecture's "
        "rigidity. Dark energy receives its operational reading as the "
        "isotropic monopole-decay pressure from the vacuum's continuous "
        "self-disruption attempts; the sign Λ > 0 is forced by the "
        "topological repulsion of monopoles from the bipartite vacuum. The "
        "298-fold amplification factor 2 × dim(E₆) × h / 2π in the Λ formula "
        "makes the universe's accelerating expansion a precision spectrometer "
        "for the substrate's zero-point Berry phase γ_Berry — the cosmic "
        "expansion rate is the substrate signaling its own zero-point "
        "geometry out at the cosmological scale."
    )

    # Find the §2.8-dynamic paragraph in §10
    anchor = None
    for p in doc.paragraphs:
        if (p.text.strip().startswith("At the dynamic level")
                and "Fano-cycling generators" in p.text):
            anchor = p
            break

    if anchor is None:
        print("ABORT: §2.8 dynamic paragraph not found in §10")
        return

    print("Found §2.8 dynamic paragraph anchor")
    new_p = make_para(new_para_text)
    anchor._element.addnext(new_p)
    print("Inserted §2.9 bipartite-vacuum paragraph in §10")

    # ----------------------------------------------------------------
    # STEP 2: Update the crown-jewel paragraph
    # ----------------------------------------------------------------
    # Find the crown jewel paragraph (starts with "The crown jewel:" or similar)
    crown_jewel_anchor = None
    for p in doc.paragraphs:
        t = p.text.strip()
        if (t.startswith("The crown jewel:")
                and "substrate's outputs are the Riemann zeros" in t):
            crown_jewel_anchor = p
            break

    if crown_jewel_anchor is None:
        print("WARN: crown jewel paragraph not found — skipping")
    else:
        new_crown_jewel = (
            "The crown jewel: the substrate's outputs are the Riemann zeros; "
            "the Riemann zeros are the observable matter quanta; the matter "
            "quanta organise into the Standard Model because the merkabit "
            "substrate IS the Fano plane PG(2,2) with the σ-fixed R-vertex "
            "(Fano 010) as its perpendicular Cartan centre, rotating under "
            "the Z_7 Fano cycling that lives in the boundary tunnel — and "
            "sitting inside a bipartite vacuum that resists every disruption. "
            "The operating plane carries the SU(2) Pauli structure on each "
            "line, the SU(3) Weyl structure on the three operating lines "
            "through R, and the E_6 positive roots as the cycling generators. "
            "The architecture produces two architectural rates simultaneously: "
            "matter creation at η_B = 5α⁴/24 (the build-success rate at the "
            "operating site) and vacuum monopole formation at Λ = √(3/2) × "
            "exp(−2γ_Berry × 78 × 12 / 2π) (the zero-point self-disruption "
            "rate at every node). The first rate populates the observable "
            "universe with matter; the second drives dark energy as isotropic "
            "monopole-decay pressure. The substrate clock advances along the "
            "x-axis (the all-T Fano line); matter emerges on the y-axis (the "
            "all-S Fano line at Re(s) = 1/2); R is the centre both "
            "perpendicular axes are arranged around. The Standard Model gauge "
            "group is the algebraic structure that this rotating projective-"
            "geometric machine carries; matter is the cycle that centres on "
            "R; antimatter is the cycle that doesn't; the Riemann zeros are "
            "the timestamps of the centring events. The Standard Model is "
            "the Fifth Face of one group acting on one lattice, on one Fano "
            "plane, with one rotation centre, two perpendicular operating "
            "axes, and one bipartite vacuum that resists having anything "
            "else written on it. The universe is the matter generator's log, "
            "the cosmological constant is its zero-point Berry-phase "
            "spectrometer reading at the cosmic scale, and the machine has "
            "only one centre on which to write."
        )
        replace_paragraph_text(crown_jewel_anchor, new_crown_jewel)
        print("Updated crown-jewel paragraph with v9 content")

    # ----------------------------------------------------------------
    # STEP 3: Update the version footer
    # ----------------------------------------------------------------
    # Find the "Paper 41, draft v6" header line
    version_header = None
    for p in doc.paragraphs:
        if "Paper 41, draft v6" in p.text:
            version_header = p
            break

    if version_header is None:
        print("WARN: version header not found — skipping footer updates")
    else:
        # Update to v9
        replace_paragraph_text(
            version_header,
            "Paper 41, draft v9. Date: 2026-05-14. Author: Stenberg, S. with Claude Anthropic."
        )

        # Insert new version entries (v9, v8, v7) after the version header,
        # before the existing v6 changelog
        v9_entry = make_para(
            "v9 changes from v8: §2.9 added (the bipartite vacuum and the Λ "
            "rate — the vacuum's own self-disruption attempts, with four "
            "subsections: §2.9.1 bipartite vacuum as global ground state; "
            "§2.9.2 three monopole types as three failure modes of the "
            "dual-spinor; §2.9.3 the Λ formula and the 298-fold Berry-phase "
            "amplification; §2.9.4 two architectural rates η_B and Λ, dark "
            "energy as isotropic monopole-decay pressure, sign Λ > 0 forced "
            "by topological repulsion). Material drawn from Paper 23 [23] "
            "(Cosmological Constant from Vacuum Monopole Suppression). §10 "
            "Conclusion gains a sixth architectural layer with the §2.9 "
            "bipartite-vacuum / two-rate / dark-energy reading; the crown-"
            "jewel statement is sharpened to include the bipartite vacuum "
            "and the two simultaneous rates."
        )
        v8_entry = make_para(
            "v8 changes from v7: structural-edit pass; §2.8.4 wording "
            "refined; consolidation of the four-layer summary into a single "
            "table."
        )
        v7_entry = make_para(
            "v7 changes from v6: §2.9 (dual-spinor / CPT) drafted in a "
            "branch; subsequent restructure folded the dual-spinor content "
            "into §2.8 and §2.9 was re-claimed for the bipartite-vacuum / Λ "
            "rate addition in v9."
        )

        # Insert in reverse order so they appear v9, v8, v7 above the v6 entry
        version_header._element.addnext(v7_entry)
        version_header._element.addnext(v8_entry)
        version_header._element.addnext(v9_entry)
        print("Inserted v7, v8, v9 changelog entries after version header")

    # Update the "four complementary" → "five complementary" summary line
    for p in doc.paragraphs:
        if "four complementary architectural descriptions" in p.text:
            full = "".join(r.text for r in p.runs)
            new = full.replace(
                "four complementary architectural descriptions",
                "five complementary architectural descriptions")
            new = new.replace(
                "§2.5 operating axes (R, T, S as rotation/imaginary/real); "
                "§2.6 Lie-algebraic (Pauli + Weyl + gauge group); §2.7 p",
                "§2.5 operating axes (R, T, S as rotation/imaginary/real); "
                "§2.6 Lie-algebraic (Pauli + Weyl + gauge group); §2.7 "
                "projective-geometric (Fano + automorphisms); §2.8 dynamic "
                "(Z_7 cycling = E_6 generation + matter/antimatter); §2.9 "
                "bipartite vacuum + Λ rate (two architectural rates, dark "
                "energy from monopole decay). P"
            )
            replace_paragraph_text(p, new)
            print("Updated summary: four-layer → five-layer description")
            break

    doc.save(str(DOC))
    print()
    print("=" * 60)
    print("DONE")
    print("=" * 60)
    print("Saved: %s" % DOC.name)


if __name__ == "__main__":
    main()
