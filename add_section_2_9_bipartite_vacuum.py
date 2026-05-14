"""
Add §2.9 'The bipartite vacuum and the Λ rate' to Paper_41_Matter_Generator_v9.docx.

Material from Paper 23 v2 (Cosmological Constant from Vacuum Monopole Suppression).
Inserts a new §2.9 after the "Summary of the four-layer architectural reading"
in §2.8, and updates that summary to reflect five layers.

Conservative: only inserts new content; doesn't touch existing prose.
Saves in place to Paper_41_Matter_Generator_v9.docx.
"""
import copy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

HERE = Path(__file__).resolve().parent
DOC = HERE / "Paper_41_Matter_Generator_v9.docx"


def make_heading(text, level=3):
    """Create a heading paragraph element."""
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    pStyle = OxmlElement("w:pStyle")
    pStyle.set(qn("w:val"), "Heading%d" % level)
    pPr.append(pStyle)
    p.append(pPr)
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    r.append(t)
    p.append(r)
    return p


def make_para(text, runs=None):
    """Create a paragraph element with text or run-spec.
    runs is a list of (text, italic, bold) tuples for mixed formatting."""
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    p.append(pPr)
    if runs is None:
        runs = [(text, False, False)]
    for txt, italic, bold in runs:
        r = OxmlElement("w:r")
        if italic or bold:
            rPr = OxmlElement("w:rPr")
            if italic:
                rPr.append(OxmlElement("w:i"))
            if bold:
                rPr.append(OxmlElement("w:b"))
            r.append(rPr)
        t = OxmlElement("w:t")
        t.text = txt
        t.set(qn("xml:space"), "preserve")
        r.append(t)
        p.append(r)
    return p


def make_list_item(text):
    """Create a bullet-point list paragraph."""
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    pStyle = OxmlElement("w:pStyle")
    pStyle.set(qn("w:val"), "ListParagraph")
    pPr.append(pStyle)
    p.append(pPr)
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "• " + text
    t.set(qn("xml:space"), "preserve")
    r.append(t)
    p.append(r)
    return p


def main():
    print("Opening %s..." % DOC.name)
    doc = Document(str(DOC))

    # Find the "Summary of the four-layer architectural reading" heading
    target_anchor = None
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if "Summary of the four-layer architectural reading" in t:
            target_anchor = p
            print("Found summary heading at paragraph %d" % i)
            break

    if target_anchor is None:
        # Fallback: find §3 heading
        for p in doc.paragraphs:
            try:
                sn = p.style.name if p.style else None
            except:
                sn = None
            if sn == "Heading 2" and p.text.strip().startswith("3."):
                target_anchor = p
                print("Fallback: inserting before §3 at: %s" % p.text[:60])
                break

    if target_anchor is None:
        print("ABORT: could not find insertion point")
        return

    # ----------------------------------------------------------------
    # Compose §2.9 content
    # ----------------------------------------------------------------
    elements_to_insert = []

    # Heading
    elements_to_insert.append(make_heading(
        "2.9 The bipartite vacuum and the Λ rate — the vacuum's own self-disruption attempts",
        level=3
    ))

    # Intro
    elements_to_insert.append(make_para(
        "§§2.5–2.8 describe the operating geometry at the active site where the "
        "12 operations of §3 fire. This section establishes what the operating "
        "geometry sits in: the bipartite vacuum that fills every node of the "
        "Eisenstein lattice when no operation is firing, and the second "
        "architectural rate — Λ, the cosmological constant — that arises from "
        "the vacuum's continuous suppression of its own self-disruption "
        "attempts. The substrate has not one but two architectural rates: η_B "
        "for matter creation at the operating site (§2.8.3), and Λ for vacuum "
        "monopole formation at every node simultaneously (this section). Both "
        "are exponential suppressions arising from different combinations of "
        "the same architectural ingredients. Paper 23 [23] derives Λ from "
        "first principles to 0.2% agreement with the observed Planck-2018 "
        "value with zero free parameters."
    ))

    # §2.9.1
    elements_to_insert.append(make_heading(
        "2.9.1 The bipartite vacuum as the global ground state",
        level=4
    ))
    elements_to_insert.append(make_para(
        "Every node of the Eisenstein lattice carries a dual-spinor "
        "(|u⟩, |v⟩) ∈ S³ × S³ in the bipartite ground state |u⟩ ⊥ |v⟩ — "
        "the zero-point standing wave (Paper 30 [30] §4.1; Paper 23 [23] "
        "§2.1). This bipartite structure is the geometric expression of the "
        "vacuum's torsion balance: forward and backward winding in perfect "
        "equilibrium at every point. The 12 operations of §3 act at a "
        "specific operating site; the bipartite structure fills every other "
        "node simultaneously as the unactivated ground state. The vacuum is "
        "rigid in all 78 directions of dim(E₆): disrupting it at any node "
        "costs energy proportional to dim(E₆) × h."
    ))

    # §2.9.2
    elements_to_insert.append(make_heading(
        "2.9.2 Three monopole types as three failure modes of the dual-spinor",
        level=4
    ))
    elements_to_insert.append(make_para(
        "Paper 23 [23] §2.1 names three architecturally distinct ways the "
        "bipartite structure can be disrupted — three monopole types — each "
        "of which the vacuum rejects within 1–4 Coxeter cycles:"
    ))
    elements_to_insert.append(make_list_item(
        "Type A: |v⟩ = 0 (missing inverse spinor). The bipartite collapses "
        "to a single spinor; CPT partner absent. Shortest lifetime; dominant "
        "contribution to Λ."
    ))
    elements_to_insert.append(make_list_item(
        "Type B: |v⟩ = |u⟩ (locked spinors). The dual structure fuses into a "
        "single state; CPT closure fails. Intermediate lifetime."
    ))
    elements_to_insert.append(make_list_item(
        "Type C: h' ≠ 12 (incommensurable winding). The cycling frequency is "
        "wrong; fails E₆ Coxeter self-duality. Longest lifetime; corrects via "
        "beat-frequency at 1/lcm(7, 12) = 1/84 cycles."
    ))
    elements_to_insert.append(make_para(
        "Types A and B are repulsive to bipartite matter (probe drift +0.10 "
        "to +0.12 lattice units over 200 cycles) and emit isotropic l = 0 "
        "monopole radiation on decay. Type C is near-transparent due to the "
        "84-cycle beat. The three failure modes correspond architecturally "
        "to three structural disruption mechanisms of the dual-spinor "
        "bipartite balance — they are not three separate processes but three "
        "ways one process (vacuum self-disruption) can fail."
    ))

    # §2.9.3
    elements_to_insert.append(make_heading(
        "2.9.3 The Λ formula — vacuum monopole formation rate",
        level=4
    ))
    elements_to_insert.append(make_para(
        "The probability of spontaneous monopole formation per node per "
        "Coxeter cycle is the Boltzmann suppression of the disruption "
        "energy. Paper 23 [23] §3 derives:"
    ))
    elements_to_insert.append(make_para(
        "Λ = √(3/2) × exp(−2 × γ_Berry × dim(E₆) × h / 2π) "
        "= 1.2247 × exp(−280.1) = 2.876 × 10⁻¹²²",
        runs=[(
            "Λ = √(3/2) × exp(−2 × γ_Berry × dim(E₆) × h / 2π) "
            "= 1.2247 × exp(−280.1) = 2.876 × 10⁻¹²²", True, False
        )]
    ))
    elements_to_insert.append(make_para(
        "Observed value (Planck 2018): Λ = 2.87 × 10⁻¹²². Ratio "
        "Λ_derived / Λ_observed = 1.002 (0.2% match). The √(3/2) factor is "
        "the ternary-binary coupling Z₃/Z₂ (Paper 24, Simulation 11). The "
        "three architectural inputs are γ_Berry = 0.94 rad (E₆ Coxeter "
        "zero-point Berry phase, Papers 1–2), dim(E₆) = 78 (McKay "
        "correspondence from P₂₄), and h = 12 (E₆ Coxeter number from "
        "Langlands self-duality). Zero free parameters."
    ))
    elements_to_insert.append(make_para(
        "The exponent factor 2 × dim(E₆) × h / (2π) = 297.94 ≈ 298 is an "
        "amplification factor: a change of δγ = 6.8 × 10⁻⁶ in γ_Berry "
        "produces a 0.2% change in Λ. The cosmological constant is a "
        "298-fold amplified measurement of the vacuum's zero-point Berry "
        "phase. The universe's accelerating expansion is the substrate "
        "signaling γ_Berry out at the cosmological scale — a precision "
        "spectrometer for the architecture's zero-point geometry."
    ))

    # §2.9.4
    elements_to_insert.append(make_heading(
        "2.9.4 Two architectural rates: η_B and Λ — dark energy as monopole-decay pressure",
        level=4
    ))
    elements_to_insert.append(make_para(
        "The architecture produces two simultaneous rates from two different "
        "combinations of the same fundamental quantities:"
    ))
    elements_to_insert.append(make_list_item(
        "η_B = 5α⁴/|2T| = 5/(24 · 137⁴) ≈ 5.91 × 10⁻¹⁰ per cycle — the "
        "matter creation success rate at the operating site (§2.8.3). "
        "Ingredients: 5 (ouroboros gates), α⁴ (four σ-equivariant gates each "
        "at α coupling), |2T| = 24 (binary tetrahedral / Fano-line stabiliser)."
    ))
    elements_to_insert.append(make_list_item(
        "Λ = √(3/2) × exp(−2γ_Berry × dim(E₆) × h / 2π) ≈ 2.87 × 10⁻¹²² per "
        "cycle per node — the vacuum monopole formation rate at every node "
        "simultaneously. Ingredients: γ_Berry (E₆ zero-point Berry phase), "
        "dim(E₆) = 78, h = 12, √(3/2) (Z₃/Z₂ ternary-binary coupling)."
    ))
    elements_to_insert.append(make_para(
        "The ratio η_B / Λ ≈ 2.6 × 10¹¹² — matter creation is exponentially "
        "more likely than vacuum self-disruption. Both are exponential "
        "suppressions from the architecture's rigidity; they use different "
        "ingredient combinations because they suppress different processes. "
        "Together they account for two of the universe's most precisely "
        "measured numbers: the cosmic baryon-to-photon ratio (η_B) and the "
        "cosmological constant (Λ)."
    ))
    elements_to_insert.append(make_para(
        "Dark energy receives its operational reading: the cosmological "
        "constant is the isotropic (l = 0) pressure from monopole formation "
        "and decay at every node simultaneously. Each lattice node "
        "spontaneously attempts a monopole with probability Λ per Coxeter "
        "cycle; the monopole decays within 1–4 cycles, emitting isotropic "
        "radiation that exerts uniform outward pressure on surrounding "
        "matter. Integrated over all nodes at all times, this pressure is Λ. "
        "Dark energy is not a field, not a particle, not a modification of "
        "gravity — it is the statistical consequence of the bipartite "
        "vacuum's intolerance of monopole disruptions, expressed as uniform "
        "outward pressure from their continuous decay. The sign Λ > 0 is "
        "forced: monopole decay is repulsive (topological incompatibility "
        "with the bipartite vacuum), so the integrated pressure accelerates "
        "expansion rather than decelerating it."
    ))

    # ----------------------------------------------------------------
    # Insert all new elements before target_anchor
    # ----------------------------------------------------------------
    target_elem = target_anchor._element
    parent = target_elem.getparent()
    target_idx = list(parent).index(target_elem)
    for i, elem in enumerate(elements_to_insert):
        parent.insert(target_idx + i, elem)
    print("Inserted %d new paragraphs (§2.9 with 4 subsections + content)"
          % len(elements_to_insert))

    # ----------------------------------------------------------------
    # Update "Summary of the four-layer..." to "five-layer"
    # ----------------------------------------------------------------
    for p in doc.paragraphs:
        t = p.text
        if "Summary of the four-layer architectural reading" in t:
            for r in p.runs:
                if "four-layer" in r.text:
                    r.text = r.text.replace("four-layer", "five-layer")
                    print("Updated summary heading: four-layer → five-layer")
                    break
            else:
                # Fallback: collapse and edit
                full = "".join(rr.text for rr in p.runs)
                new = full.replace("four-layer", "five-layer")
                for rr in list(p.runs)[1:]:
                    rr._element.getparent().remove(rr._element)
                if p.runs:
                    p.runs[0].text = new
                else:
                    p.add_run(new)
                print("Updated summary heading (fallback): four-layer → five-layer")
            break

    doc.save(str(DOC))
    print()
    print("=" * 60)
    print("DONE")
    print("=" * 60)
    print("Saved: %s" % DOC.name)
    print()
    print("v9 now has §2.9 with four subsections:")
    print("  2.9.1 The bipartite vacuum as the global ground state")
    print("  2.9.2 Three monopole types as three failure modes")
    print("  2.9.3 The Λ formula — vacuum monopole formation rate")
    print("  2.9.4 Two architectural rates: η_B and Λ")


if __name__ == "__main__":
    main()
